"""
Generación de informes a partir de plantillas Word (.docx) del usuario.

El usuario sube su propia plantilla con marcadores estilo Jinja2:

    Obra: {{ obra.nombre }}
    Cliente: {{ obra.cliente }}          Avance: {{ kpi.avance }} %

Para repetir filas de una tabla de Word hacen falta tres filas: una con la
apertura del bucle, otra con los datos y otra con el cierre. Las de control
desaparecen al renderizar:

    ┌──────────────────────────────┐
    │ {%tr for x in dispositivos %}│
    ├──────────┬─────────┬─────────┤
    │{{x.etiqueta}}│{{x.modelo}}│{{x.ip}}│
    ├──────────┴─────────┴─────────┤
    │ {%tr endfor %}               │
    └──────────────────────────────┘

Poner la apertura y el cierre dentro de la misma fila —lo que sugiere la
documentación de docxtpl— falla en la versión 0.20 con «unknown tag endfor».

Se rellena con los datos reales de la obra y se descarga el .docx resultante,
conservando membrete, estilos, tablas y pies de página de la plantilla original.

Si `docxtpl` no está instalado se usa un modo de sustitución simple que sólo
reemplaza {{ campo }} sueltos (sin bucles), para que nunca falle del todo.
"""
from __future__ import annotations

import io
import re
from datetime import date, datetime
from pathlib import Path
from typing import Any

from .. import db
from . import kpis


def _fmt_fecha(valor) -> str:
    if not valor:
        return ""
    try:
        return datetime.fromisoformat(str(valor)[:10]).strftime("%d/%m/%Y")
    except ValueError:
        return str(valor)


def _fmt_eur(valor) -> str:
    try:
        n = float(valor or 0)
    except (TypeError, ValueError):
        return str(valor or "")
    return f"{n:,.2f} €".replace(",", " ").replace(".", ",").replace(" ", ".")


def _resolver_refs(key: str, filas: list[dict]) -> list[dict]:
    """Sustituye los *_id por su etiqueta legible, para que la plantilla sea limpia."""
    from ..schema import ENTITIES
    ent = ENTITIES.get(key)
    if not ent:
        return filas
    refs = {f.name: f.ref for f in ent.fields if f.type == "ref" and f.ref}
    if not refs:
        return filas
    cache: dict[tuple[str, int], str] = {}
    out = []
    for fila in filas:
        f = dict(fila)
        for campo, destino in refs.items():
            rid = f.get(campo)
            if not rid:
                f[campo.replace("_id", "")] = ""
                continue
            clave = (destino, int(rid))
            if clave not in cache:
                reg = db.obtener(destino, int(rid)) or {}
                cache[clave] = str(reg.get(ENTITIES[destino].title_field, "") or "")
            f[campo.replace("_id", "")] = cache[clave]
        out.append(f)
    return out


def contexto(obra_id: int) -> dict[str, Any]:
    """Todo lo que una plantilla puede usar."""
    r = kpis.resumen(obra_id)
    obra = dict(r["obra"])

    # Versiones formateadas de las fechas e importes de la obra.
    for campo in list(obra):
        if campo.startswith("fecha_"):
            obra[campo + "_txt"] = _fmt_fecha(obra[campo])
    obra["importe_contrato_txt"] = _fmt_eur(obra.get("importe_contrato"))

    ctx: dict[str, Any] = {
        "obra": obra,
        "hoy": date.today().strftime("%d/%m/%Y"),
        "hoy_iso": date.today().isoformat(),
        "ahora": datetime.now().strftime("%d/%m/%Y %H:%M"),
        "kpi": {
            "avance": r["tareas"]["avance"],
            "tareas_total": r["tareas"]["total"],
            "tareas_completadas": r["tareas"]["completadas"],
            "tareas_retrasadas": r["tareas"]["retrasadas"],
            "tareas_bloqueadas": r["tareas"]["bloqueadas"],
            "dispositivos_total": r["dispositivos"]["total"],
            "dispositivos_unidades": r["dispositivos"]["unidades"],
            "dispositivos_instalados": r["dispositivos"]["instalados"],
            "dispositivos_probados": r["dispositivos"]["probados"],
            "pct_instalado": r["dispositivos"]["pct_instalado"],
            "incidencias_abiertas": r["otros"]["incidencias_abiertas"],
            "docs_pendientes": r["otros"]["docs_pendientes"],
            "personal": r["otros"]["personal"],
            "dias_restantes": r["plazo"]["dias_restantes"],
        },
        "economico": {k: v for k, v in r["economico"].items() if k != "desglose"},
        "economico_txt": {
            k: _fmt_eur(v) for k, v in r["economico"].items()
            if isinstance(v, (int, float))
        },
        "alertas": r["alertas"],
        "stock": kpis.stock(obra_id),
    }

    # Todas las tablas de la obra, con referencias resueltas.
    from ..schema import ENTITIES
    for key, ent in ENTITIES.items():
        if key == "obras":
            continue
        filas = db.listar(key, obra_id if ent.per_obra else None)
        ctx[key] = _resolver_refs(key, filas)

    # Subconjuntos de uso frecuente en actas e informes.
    ctx["tareas_abiertas"] = [t for t in ctx["tareas"] if t.get("estado") not in ("Completada", "Cancelada")]
    ctx["tareas_completadas"] = [t for t in ctx["tareas"] if t.get("estado") == "Completada"]
    ctx["incidencias_abiertas"] = [i for i in ctx["incidencias"] if i.get("estado") not in ("Resuelta", "Cerrada")]
    ctx["dispositivos_instalados"] = [
        d for d in ctx["dispositivos"]
        if d.get("estado") in ("Instalado", "Conexionado", "Configurado", "Probado", "Entregado")
    ]
    ctx["camaras"] = [d for d in ctx["dispositivos"] if d.get("categoria") == "CCTV"]
    ctx["material_reponer"] = [s for s in ctx["stock"] if s.get("alerta") != "OK"]
    return ctx


# ─────────────────────────────────────────────────── generación con docxtpl
def generar_docx(plantilla: Path, obra_id: int) -> bytes:
    ctx = contexto(obra_id)
    try:
        from docxtpl import DocxTemplate
    except ImportError:
        return _generar_docx_simple(plantilla, ctx)

    doc = DocxTemplate(str(plantilla))
    entorno_filtros = {"eur": _fmt_eur, "fecha": _fmt_fecha}
    try:
        import jinja2
        env = jinja2.Environment()
        env.filters.update(entorno_filtros)
        doc.render(ctx, env)
    except ImportError:
        doc.render(ctx)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


_MARCADOR = re.compile(r"\{\{\s*([a-zA-Z0-9_.]+)\s*\}\}")


def _valor(ctx: dict, ruta: str) -> str:
    actual: Any = ctx
    for parte in ruta.split("."):
        if isinstance(actual, dict):
            actual = actual.get(parte)
        else:
            actual = getattr(actual, parte, None)
        if actual is None:
            return ""
    return str(actual)


def _generar_docx_simple(plantilla: Path, ctx: dict) -> bytes:
    """Sustitución de {{ campo }} sin bucles. Reserva si falta docxtpl."""
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(
            "Para generar informes Word instala: pip install docxtpl python-docx"
        ) from exc

    doc = Document(str(plantilla))

    def sustituir_parrafo(p):
        texto = "".join(run.text for run in p.runs)
        if "{{" not in texto:
            return
        nuevo = _MARCADOR.sub(lambda m: _valor(ctx, m.group(1)), texto)
        if nuevo == texto:
            return
        for run in p.runs[1:]:
            run.text = ""
        if p.runs:
            p.runs[0].text = nuevo

    for p in doc.paragraphs:
        sustituir_parrafo(p)
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for p in celda.paragraphs:
                    sustituir_parrafo(p)
    for seccion in doc.sections:
        for contenedor in (seccion.header, seccion.footer):
            for p in contenedor.paragraphs:
                sustituir_parrafo(p)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ────────────────────────────────────────── informe por defecto (sin plantilla)
def informe_estandar(obra_id: int) -> bytes:
    """Informe de seguimiento completo, por si el usuario aún no ha subido plantilla."""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt, RGBColor
    except ImportError as exc:
        raise RuntimeError("Instala python-docx: pip install python-docx") from exc

    ctx = contexto(obra_id)
    obra, kpi, eco = ctx["obra"], ctx["kpi"], ctx["economico"]

    doc = Document()
    est = doc.styles["Normal"]
    est.font.name = "Calibri"
    est.font.size = Pt(10)

    t = doc.add_heading("INFORME DE SEGUIMIENTO DE OBRA", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(f"{obra.get('nombre', '')} · {ctx['ahora']}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def tabla_pares(titulo, pares):
        doc.add_heading(titulo, level=1)
        tb = doc.add_table(rows=0, cols=2)
        tb.style = "Light Grid Accent 1"
        for k, v in pares:
            if v in (None, "", 0):
                continue
            fila = tb.add_row().cells
            fila[0].text = str(k)
            fila[1].text = str(v)
        doc.add_paragraph()

    tabla_pares("1. Datos del proyecto", [
        ("Proyecto", obra.get("nombre")), ("Código / OT", obra.get("codigo")),
        ("Cliente", obra.get("cliente")), ("Cliente final", obra.get("cliente_final")),
        ("Emplazamiento", f"{obra.get('direccion') or ''} {obra.get('poblacion') or ''}".strip()),
        ("Jefe de obra", obra.get("jefe_obra")),
        ("Interlocutor cliente", obra.get("interlocutor_cliente")),
        ("Estado", obra.get("estado")),
        ("Inicio", _fmt_fecha(obra.get("fecha_inicio"))),
        ("Fin previsto", _fmt_fecha(obra.get("fecha_fin_prevista"))),
        ("Importe contrato", _fmt_eur(obra.get("importe_contrato"))),
    ])

    tabla_pares("2. Situación general", [
        ("Avance global", f"{kpi['avance']} %"),
        ("Tareas", f"{kpi['tareas_completadas']} completadas de {kpi['tareas_total']}"),
        ("Tareas con retraso", kpi["tareas_retrasadas"]),
        ("Tareas bloqueadas", kpi["tareas_bloqueadas"]),
        ("Días restantes de plazo", kpi["dias_restantes"]),
        ("Dispositivos instalados", f"{kpi['dispositivos_instalados']} de {kpi['dispositivos_total']} ({kpi['pct_instalado']} %)"),
        ("Dispositivos probados", kpi["dispositivos_probados"]),
        ("Incidencias abiertas", kpi["incidencias_abiertas"]),
        ("Documentos pendientes", kpi["docs_pendientes"]),
    ])

    tabla_pares("3. Situación económica", [
        ("Importe de contrato", _fmt_eur(eco["contrato"])),
        ("Ampliaciones aprobadas", _fmt_eur(eco["ampliaciones"])),
        ("Ingreso total", _fmt_eur(eco["ingresos"])),
        ("Coste incurrido", _fmt_eur(eco["coste_total"])),
        ("Margen bruto", _fmt_eur(eco["margen"])),
        ("Margen", f"{eco['margen_pct']} %"),
        ("Certificado", _fmt_eur(eco["certificado"])),
        ("Pendiente de cobro", _fmt_eur(eco["pendiente_cobro"])),
    ])

    if ctx["alertas"]:
        doc.add_heading("4. Puntos de atención", level=1)
        for a in ctx["alertas"][:15]:
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(f"[{a['severidad'].upper()}] {a['titulo']}. ")
            r.bold = True
            if a["severidad"] == "critica":
                r.font.color.rgb = RGBColor(0xC0, 0x1A, 0x1A)
            p.add_run(a["detalle"])
        doc.add_paragraph()

    def tabla_datos(titulo, filas, columnas):
        if not filas:
            return
        doc.add_heading(titulo, level=1)
        tb = doc.add_table(rows=1, cols=len(columnas))
        tb.style = "Light Grid Accent 1"
        for i, (_, cab) in enumerate(columnas):
            celda = tb.rows[0].cells[i]
            celda.text = cab
            for p in celda.paragraphs:
                for run in p.runs:
                    run.bold = True
        for fila in filas:
            celdas = tb.add_row().cells
            for i, (campo, _) in enumerate(columnas):
                v = fila.get(campo)
                if campo.startswith("fecha"):
                    v = _fmt_fecha(v)
                celdas[i].text = "" if v in (None, "") else str(v)
        doc.add_paragraph()

    tabla_datos("5. Tareas en curso y pendientes", ctx["tareas_abiertas"][:40], [
        ("tarea", "Tarea"), ("categoria", "Categoría"), ("estado", "Estado"),
        ("avance", "%"), ("fecha_fin", "Fin previsto"), ("responsable", "Responsable"),
    ])
    tabla_datos("6. Incidencias abiertas", ctx["incidencias_abiertas"][:25], [
        ("titulo", "Incidencia"), ("gravedad", "Gravedad"), ("estado", "Estado"),
        ("responsable", "Responsable"), ("fecha_limite", "Límite"),
    ])
    tabla_datos("7. Inventario de instalación", ctx["dispositivos"][:120], [
        ("etiqueta", "Etiqueta"), ("tipo", "Tipo"), ("marca", "Marca"),
        ("modelo", "Modelo"), ("num_serie", "Nº serie"), ("ip", "IP"),
        ("zona", "Zona"), ("estado", "Estado"),
    ])
    tabla_datos("8. Material bajo mínimo", ctx["material_reponer"][:25], [
        ("material", "Material"), ("restante", "Restante"), ("stock_min", "Mínimo"),
        ("alerta", "Alerta"), ("proveedor", "Proveedor"),
    ])

    doc.add_paragraph()
    doc.add_paragraph("Documento generado automáticamente por ObraSec.").italic = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def campos_disponibles(obra_id: int | None = None) -> dict:
    """Lista de marcadores que el usuario puede usar en sus plantillas."""
    from ..schema import ENTITIES
    salida = {
        "Datos de la obra": [f"obra.{f.name}" for f in ENTITIES["obras"].fields],
        "Generales": ["hoy", "ahora", "hoy_iso"],
        "Indicadores": [
            "kpi.avance", "kpi.tareas_total", "kpi.tareas_completadas",
            "kpi.tareas_retrasadas", "kpi.tareas_bloqueadas", "kpi.dispositivos_total",
            "kpi.dispositivos_instalados", "kpi.dispositivos_probados",
            "kpi.pct_instalado", "kpi.incidencias_abiertas", "kpi.docs_pendientes",
            "kpi.dias_restantes",
        ],
        "Económico": [
            "economico.contrato", "economico.ampliaciones", "economico.ingresos",
            "economico.coste_total", "economico.margen", "economico.margen_pct",
            "economico.certificado", "economico.pendiente_cobro",
            "economico_txt.contrato", "economico_txt.coste_total", "economico_txt.margen",
        ],
    }
    for key, ent in ENTITIES.items():
        if key == "obras":
            continue
        salida[f"Tabla: {ent.plural}"] = [
            f"{{% for x in {key} %}} … {{{{ x.{f.name} }}}} … {{% endfor %}}"
            for f in ent.fields[:1]
        ] + [f"x.{f.name}" for f in ent.fields]
    return salida
